from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Crear el documento
document = Document()

# Primer parrafo - Centrado con diferentes estilos en el mismo parrafo
p1 = document.add_paragraph() # Variable que añade un parrafo al documento
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER # Valor de centrado en el documento
run1 = p1.add_run("Este es un texto") # Crear la variable y le añade un texto
run1.font.name = 'Arial' # tipo de letra
run1.font.size = Pt(12) # tamaño de letra
run1.font.bold = True # Negrita
run2 = p1.add_run(" y este otro texto") # Crear la variable y le añade un texto
run2.font.name = 'Times New Roman' # tipo de letra
run2.font.size = Pt(14) # tamaño de letra
run2.font.italic = True # Itálico
run3 = p1.add_run(" y este otro texto en color rojo.") # Crear la variable y le añade un texto
run3.font.name = 'Arial' # tipo de letra
run3.font.size = Pt(16) # tamaño de letra
run3.font.color.rgb = RGBColor(255,0,0) # Color rojo


# Guarda el documento
document.save("documento_Multiformato.docx")