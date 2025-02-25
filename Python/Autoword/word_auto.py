import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class WordAuto:
    def __init__(self, filename="documento_auto.docx"):
        self.filename = filename
        self.doc = Document()

    # Método para agregar un párrafo al documento
    def add_paragraph(self, text, fuente='Arial', tamano=12):
        paragraph = self.doc.add_paragraph(text)
        run = paragraph.runs[0]
        run.font.name = fuente
        run.font.size = Pt(tamano)
    
    # Método para agregar una tabla al documento
    def add_table(self, rows, cols):
        table = self.doc.add_table(rows=rows, cols=cols)
        for row in table.rows:
            for cell in row.cells:
                cell.text = "Contenido"
    
    # Método para guardar el documento
    def save_document(self):
        self.doc.save(self.filename)









 